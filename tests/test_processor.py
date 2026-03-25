"""
tests/test_processor.py — Unit tests for DocxProcessor and replacer helpers.

Run with:  python -m pytest tests/ -v
"""

import sys, os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

import tempfile
import pytest

from docx import Document
from core.replacer import replace_in_paragraph, replace_in_text
from core.processor import DocxProcessor, make_output_path


# ── replace_in_text (plain string) ───────────────────────────────────────────

class TestReplaceInText:

    def test_simple_replacement(self) -> None:
        text, n = replace_in_text("Hello World", {"World": "[NAME]"})
        assert text == "Hello [NAME]"
        assert n == 1

    def test_multiple_occurrences(self) -> None:
        text, n = replace_in_text("A B A B A", {"A": "X"})
        assert text == "X B X B X"
        assert n == 3

    def test_no_match(self) -> None:
        text, n = replace_in_text("Hello", {"Bye": "X"})
        assert text == "Hello"
        assert n == 0

    def test_longer_pattern_wins(self) -> None:
        # "ABC" should be replaced before "A" so we don't get partial replaces
        text, n = replace_in_text("ABC and A", {"ABC": "[COMPANY]", "A": "[LETTER]"})
        assert "[COMPANY]" in text
        assert n >= 1

    def test_empty_replacements(self) -> None:
        text, n = replace_in_text("hello", {})
        assert text == "hello"
        assert n == 0


# ── replace_in_paragraph (docx) ──────────────────────────────────────────────

def _make_para(text: str, bold_first: bool = False) -> "docx.text.paragraph.Paragraph":
    """Create an in-memory Document and return its first paragraph."""
    doc  = Document()
    para = doc.add_paragraph()
    run  = para.add_run(text)
    if bold_first:
        run.bold = True
    return para


class TestReplaceInParagraph:

    def test_basic_replacement(self) -> None:
        para = _make_para("The Company agrees to pay.")
        n    = replace_in_paragraph(para, {"Company": "Party A"})
        assert n == 1
        assert "Party A" in para.runs[0].text

    def test_no_change_when_no_match(self) -> None:
        para = _make_para("No sensitive info here.")
        n    = replace_in_paragraph(para, {"XYZ Corp": "[COMPANY]"})
        assert n == 0
        assert para.runs[0].text == "No sensitive info here."

    def test_multiple_replacements(self) -> None:
        para = _make_para("Call John at 13812345678.")
        n    = replace_in_paragraph(para, {
            "John": "[NAME]",
            "13812345678": "[PHONE]",
        })
        assert n == 2
        combined = "".join(r.text for r in para.runs)
        assert "[NAME]" in combined
        assert "[PHONE]" in combined

    def test_empty_para(self) -> None:
        doc  = Document()
        para = doc.add_paragraph("")
        n    = replace_in_paragraph(para, {"X": "Y"})
        assert n == 0


# ── DocxProcessor ────────────────────────────────────────────────────────────

class TestDocxProcessor:

    def _make_docx(self, content: dict[str, str]) -> str:
        """Write a temp docx with {paragraph_text: ...} and return path."""
        doc = Document()
        for text in content.values():
            doc.add_paragraph(text)
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        doc.save(tmp.name)
        tmp.close()
        return tmp.name

    def test_process_saves_file(self) -> None:
        src  = self._make_docx({"p1": "XYZ Corp Limited is the buyer."})
        out  = src.replace(".docx", "_out.docx")
        proc = DocxProcessor()
        proc.process(src, {"XYZ Corp Limited": "[COMPANY]"}, out)
        assert os.path.exists(out)
        os.unlink(src)
        os.unlink(out)

    def test_replacement_appears_in_output(self) -> None:
        src  = self._make_docx({"p1": "Contact: john@example.com"})
        out  = src.replace(".docx", "_out.docx")
        proc = DocxProcessor()
        proc.process(src, {"john@example.com": "[EMAIL]"}, out)

        result_doc = Document(out)
        texts      = [p.text for p in result_doc.paragraphs]
        assert any("[EMAIL]" in t for t in texts)
        os.unlink(src)
        os.unlink(out)

    def test_extract_text(self) -> None:
        src  = self._make_docx({"p1": "Hello World", "p2": "Second para"})
        proc = DocxProcessor()
        text = proc.extract_text(src)
        assert "Hello World" in text
        assert "Second para" in text
        os.unlink(src)


# ── DocxTextBoxes ─────────────────────────────────────────────────────────────

class TestDocxTextBoxes:
    """Verify that text inside w:txbxContent (text boxes) is extracted and replaced."""

    def _make_docx_with_textbox(self, term: str) -> str:
        """Create a temp docx with a w:txbxContent element containing *term*."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc  = Document()
        body = doc.element.body

        txbx   = OxmlElement('w:txbxContent')
        p_elem = OxmlElement('w:p')
        r_elem = OxmlElement('w:r')
        t_elem = OxmlElement('w:t')
        t_elem.text = term
        r_elem.append(t_elem)
        p_elem.append(r_elem)
        txbx.append(p_elem)
        # Insert before the last child (w:sectPr)
        body.insert(len(body) - 1, txbx)

        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        doc.save(tmp.name)
        tmp.close()
        return tmp.name

    def test_extract_text_includes_textbox(self) -> None:
        src  = self._make_docx_with_textbox("SECRET_TEXTBOX_TERM")
        proc = DocxProcessor()
        text = proc.extract_text(src)
        assert "SECRET_TEXTBOX_TERM" in text
        os.unlink(src)

    def test_process_replaces_textbox_content(self) -> None:
        from core.processor import _collect_all_paragraphs

        src  = self._make_docx_with_textbox("SECRET_TEXTBOX_TERM")
        out  = src.replace(".docx", "_out.docx")
        proc = DocxProcessor()
        proc.process(src, {"SECRET_TEXTBOX_TERM": "[REDACTED]"}, out)

        result_doc = Document(out)
        all_text   = "\n".join(p.text for p in _collect_all_paragraphs(result_doc))
        assert "[REDACTED]" in all_text
        assert "SECRET_TEXTBOX_TERM" not in all_text
        os.unlink(src)
        os.unlink(out)


# ── make_output_path ─────────────────────────────────────────────────────────

class TestMakeOutputPath:

    def test_docx_suffix(self) -> None:
        p = make_output_path("/tmp/contract.docx", "_脱敏")
        assert p == "/tmp/contract_脱敏.docx"

    def test_pdf_suffix(self) -> None:
        p = make_output_path("/tmp/agreement.pdf", "_redacted")
        assert p == "/tmp/agreement_redacted.pdf"
